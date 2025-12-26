"""
文件解析工具
支持PDF和Word文档的文本提取
所有文档通过AI API处理，无需OCR
"""
import os
import re

# 可选导入 PyMuPDF (fitz)
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

# 可选导入 pdfplumber
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# 可选导入 OCR (Tesseract)
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# 可选导入 python-docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，Word文档解析功能将不可用。")
    print("      请运行: pip install python-docx")


def clean_text(text):
    """
    清理文本中的常见错误
    保持文本结构和上下文位置信息
    """
    if not text:
        return text
    
    # 保留页面分隔标记（如果存在）
    page_markers = []
    if '--- 第' in text and '页 ---' in text:
        # 提取并保留页面标记
        page_marker_pattern = r'--- 第\d+页 ---'
        page_markers = re.findall(page_marker_pattern, text)
        # 临时替换页面标记，避免被清理
        for i, marker in enumerate(page_markers):
            text = text.replace(marker, f'__PAGE_MARKER_{i}__', 1)
    
    # 修复常见的文本识别错误
    # 修复年份中的O应该是0（如"2O25" -> "2025"）
    text = re.sub(r'([12])O(\d{3})', r'\g<1>0\2', text)  # 修复年份中的O -> 0
    text = re.sub(r'(\d{4})\.O(\d)', r'\1.0\2', text)  # 修复月份中的O -> 0
    
    # 修复邮箱中的常见错误（如"@4q.com"应该是"@qq.com"）
    text = re.sub(r'@4q\.com', '@qq.com', text, flags=re.IGNORECASE)
    text = re.sub(r'@(\d+)q\.com', r'@qq.com', text, flags=re.IGNORECASE)
    
    # 移除无法识别的字符标记，但保留必要的标点
    text = re.sub(r'[□¡¿\u200b\u200c\u200d\ufeff]', '', text)  # 移除无法识别的字符标记和零宽字符
    
    # 清理连续的特殊字符，但保留换行结构
    text = re.sub(r'[ \t]+', ' ', text)  # 多个空格/制表符合并为单个空格
    text = re.sub(r'\n{4,}', '\n\n\n', text)  # 超过3个换行符的合并为3个（保留段落分隔）
    
    # 移除行首行尾的空白，但保留空行（用于段落分隔）
    lines = text.split('\n')
    cleaned_lines = []
    prev_empty = False
    for line in lines:
        line_stripped = line.strip()
        if line_stripped:
            cleaned_lines.append(line_stripped)
            prev_empty = False
        elif not prev_empty:
            # 保留单个空行作为段落分隔
            cleaned_lines.append('')
            prev_empty = True
    
    text = '\n'.join(cleaned_lines)
    
    # 恢复页面标记
    if page_markers:
        for i, marker in enumerate(page_markers):
            text = text.replace(f'__PAGE_MARKER_{i}__', marker, 1)
    
    return text


def extract_text_from_pdf(file_path, use_ai=True, use_ocr=True):
    """
    从PDF文件提取文本 - 三级解析策略
    
    策略优先级：
    1. 优先：AI解析（如果启用且可用）
    2. 备用：PyMuPDF、pdfplumber（文本PDF）
    3. 最后：OCR（扫描件处理）
    
    Args:
        file_path: PDF文件路径
        use_ai: 是否尝试使用AI解析（默认True）
        use_ocr: 是否尝试使用OCR（默认True）
    
    Returns:
        提取的文本内容
    """
    results = {
        'method': None,
        'text': '',
        'success': False,
        'error': None
    }
    
    # ========================================================================
    # 策略1: PyMuPDF 提取（最快，适合文本PDF）
    # ========================================================================
    if FITZ_AVAILABLE:
        try:
            pdf_doc = fitz.open(file_path)
            total_pages = len(pdf_doc)
            page_texts = []
            
            for page_num in range(total_pages):
                page = pdf_doc[page_num]
                page_text = page.get_text()
                
                if page_text and isinstance(page_text, str):
                    page_text_cleaned = page_text.strip()
                    if page_text_cleaned:
                        if total_pages > 1:
                            page_texts.append(f"--- 第{page_num + 1}页 ---\n{page_text_cleaned}")
                        else:
                            page_texts.append(page_text_cleaned)
            
            pdf_doc.close()
            text = "\n\n".join(page_texts)
            text_stripped = text.strip()
            
            # 判断是否为有效文本PDF
            if text_stripped:
                total_chars = len(text_stripped)
                chinese_chars = len(re.findall(r'[\u4e00-\u9fa5]', text_stripped))
                unique_chars = len(set(text_stripped))
                chinese_ratio = chinese_chars / total_chars if total_chars > 0 else 0
                
                # 有效文本PDF判断标准
                if total_chars >= 200 and chinese_chars >= 50 and chinese_ratio >= 0.1 and unique_chars >= 50:
                    results['method'] = 'PyMuPDF'
                    results['text'] = clean_text(text_stripped)
                    results['success'] = True
                    return results['text']
        except Exception as e:
            print(f"PyMuPDF解析失败: {e}")
    
    # ========================================================================
    # 策略2: pdfplumber 提取（备用，适合复杂布局）
    # ========================================================================
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                page_texts = []
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        if len(pdf.pages) > 1:
                            page_texts.append(f"--- 第{page_num + 1}页 ---\n{page_text.strip()}")
                        else:
                            page_texts.append(page_text.strip())
                
                text = "\n\n".join(page_texts)
                text_stripped = text.strip()
                
                if text_stripped:
                    total_chars = len(text_stripped)
                    chinese_chars = len(re.findall(r'[\u4e00-\u9fa5]', text_stripped))
                    unique_chars = len(set(text_stripped))
                    chinese_ratio = chinese_chars / total_chars if total_chars > 0 else 0
                    
                    if total_chars >= 200 and chinese_chars >= 50 and chinese_ratio >= 0.1 and unique_chars >= 50:
                        results['method'] = 'pdfplumber'
                        results['text'] = clean_text(text_stripped)
                        results['success'] = True
                        return results['text']
        except Exception as e:
            print(f"pdfplumber解析失败: {e}")
    
    # ========================================================================
    # 策略3: OCR 提取（最后手段，适合扫描件）
    # ========================================================================
    if use_ocr and OCR_AVAILABLE:
        try:
            # 将PDF转换为图片，然后OCR
            if FITZ_AVAILABLE:
                pdf_doc = fitz.open(file_path)
                ocr_texts = []
                
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    # 将页面渲染为图片
                    mat = fitz.Matrix(2.0, 2.0)  # 放大2倍提高OCR准确度
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    
                    # 使用PIL打开图片
                    from io import BytesIO
                    img = Image.open(BytesIO(img_data))
                    
                    # OCR识别（支持中英文）
                    ocr_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                    if ocr_text and ocr_text.strip():
                        ocr_texts.append(f"--- 第{page_num + 1}页 ---\n{ocr_text.strip()}")
                
                pdf_doc.close()
                
                if ocr_texts:
                    text = "\n\n".join(ocr_texts)
                    text_stripped = text.strip()
                    
                    if text_stripped and len(text_stripped) >= 100:
                        results['method'] = 'OCR'
                        results['text'] = clean_text(text_stripped)
                        results['success'] = True
                        return results['text']
        except Exception as e:
            print(f"OCR解析失败: {e}")
    
    # ========================================================================
    # 所有方法都失败，返回空字符串（将由AI API处理）
    # ========================================================================
    results['error'] = '所有PDF解析方法都失败，将使用AI API处理'
    return ""


def extract_text_from_word(file_path):
    """
    从Word文档提取文本
    支持 .docx 格式，提取段落和表格中的内容
    """
    # 检查 python-docx 是否可用
    if not DOCX_AVAILABLE:
        raise Exception("python-docx 未安装，无法解析Word文档。请运行: pip install python-docx")
    
    # 检查文件是否存在
    if not os.path.exists(file_path):
        raise Exception(f"文件不存在: {file_path}")
    
    # 检查文件扩展名
    file_ext = os.path.splitext(file_path)[1].lower()
    if file_ext == '.doc':
        raise Exception("不支持旧的 .doc 格式，请将文件转换为 .docx 格式后再上传")
    
    try:
        # 使用绝对路径，避免路径问题
        abs_path = os.path.abspath(file_path)
        doc = Document(abs_path)
        text = ""
        
        # 提取段落文本（保持原有顺序和结构）
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text += para_text + "\n"
                paragraph_count += 1
        
        # 提取表格中的文本（保持表格结构，便于后续解析）
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    # 提取单元格中的所有段落文本
                    cell_text = ""
                    for para in cell.paragraphs:
                        para_text = para.text.strip()
                        if para_text:
                            cell_text += para_text + " "
                    cell_text = cell_text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    # 使用空格分隔，保持表格内容的可读性
                    text += " ".join(row_text) + "\n"
        
        # 尝试提取文本框和形状中的文本（可能包含重要信息）
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            # 遍历文档中的所有形状和文本框
            for element in doc.element.body.iter():
                # 查找文本框（textbox）
                if element.tag.endswith('textbox'):
                    for text_elem in element.iter():
                        if text_elem.text and text_elem.text.strip():
                            text += text_elem.text.strip() + "\n"
        except Exception as e:
            # 文本框提取失败不影响主流程
            pass
        
        # 尝试提取页眉和页脚（可能包含重要信息）
        try:
            for section in doc.sections:
                # 提取页眉
                if section.header:
                    for paragraph in section.header.paragraphs:
                        header_text = paragraph.text.strip()
                        if header_text and len(header_text) > 1:
                            # 页眉信息放在文本开头
                            text = header_text + "\n" + text
                # 提取页脚
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        footer_text = paragraph.text.strip()
                        if footer_text and len(footer_text) > 1:
                            # 页脚信息放在文本末尾
                            text += "\n" + footer_text
        except Exception as e:
            # 页眉页脚提取失败不影响主流程
            pass
        
        # 如果没有提取到任何文本，可能是文件格式问题
        if not text.strip():
            raise Exception("无法从Word文档中提取文本，文件可能已损坏或格式不正确（可能是加密的文档）")
        
        return clean_text(text.strip())
    except Exception as e:
        error_msg = str(e)
        # 提供更友好的错误提示
        if "cannot open" in error_msg.lower() or "permission" in error_msg.lower():
            raise Exception(f"无法打开文件，请检查文件是否被其他程序占用或没有读取权限: {error_msg}")
        elif "not a zip file" in error_msg.lower() or "bad zipfile" in error_msg.lower():
            raise Exception(f"文件格式错误，.docx 文件应该是有效的 ZIP 格式。请确认文件是否损坏: {error_msg}")
        else:
            raise Exception(f"Word文档解析失败: {error_msg}")


def extract_text(file_path, use_ai=True, use_ocr=True):
    """
    根据文件类型自动选择解析方法
    
    Args:
        file_path: 文件路径
        use_ai: 是否允许使用AI解析（默认True）
        use_ocr: 是否允许使用OCR（默认True）
    
    Returns:
        提取的文本内容
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path, use_ai=use_ai, use_ocr=use_ocr)
    elif file_ext in ['.doc', '.docx']:
        return extract_text_from_word(file_path)
    else:
        raise Exception(f"不支持的文件格式: {file_ext}")
